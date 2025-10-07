#!/usr/bin/env python3
# -*- coding: utf-8 -*-

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import mm
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.lib.enums import TA_LEFT, TA_CENTER
import os

# 日本語フォント設定（Windowsの場合）
try:
    pdfmetrics.registerFont(TTFont('MSGothic', 'msgothic.ttc'))
    pdfmetrics.registerFont(TTFont('MSMincho', 'msmincho.ttc'))
    japanese_font = 'MSGothic'
except:
    japanese_font = 'Helvetica'

def create_shinnyusyain_faq_docx():
    """新入社員FAQ DOCX作成"""
    doc = Document()

    # タイトル
    title = doc.add_heading('新入社員FAQ', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # サブタイトル
    subtitle = doc.add_paragraph('株式会社テックイノベーション 人事部')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph('最終更新: 2024年10月1日')
    doc.add_paragraph()

    # 目次
    doc.add_heading('目次', 1)
    doc.add_paragraph('1. 入社前の準備')
    doc.add_paragraph('2. 初日の流れ')
    doc.add_paragraph('3. 勤怠・給与関連')
    doc.add_paragraph('4. 福利厚生')
    doc.add_paragraph('5. 社内システム')
    doc.add_paragraph('6. オフィス環境')
    doc.add_paragraph('7. その他')
    doc.add_page_break()

    # FAQ内容
    faqs = [
        ("1. 入社前の準備", [
            ("Q1: 入社初日に必要な持ち物は何ですか？",
             "A: 以下をご持参ください。\n・身分証明書（運転免許証またはマイナンバーカード）\n・印鑑（認印）\n・通帳またはキャッシュカードのコピー（給与振込口座確認用）\n・筆記用具、ノート\n・入社書類一式（事前に郵送した書類に記入・捺印したもの）"),

            ("Q2: 服装はどうすればいいですか？",
             "A: 初日はビジネスフォーマル（スーツ）でお越しください。2日目以降はビジネスカジュアルでOKです。当社にドレスコードはなく、清潔感のある服装であれば問題ありません。"),

            ("Q3: 出社時間は何時ですか？",
             "A: 初日は9:00に1階受付にお越しください。人事担当が迎えに参ります。通常の始業時刻は9:00ですが、フレックスタイム制を導入しており、コアタイム（10:30〜15:30）に勤務していれば、出社・退社時刻は自由です。"),
        ]),

        ("2. 初日の流れ", [
            ("Q4: 初日はどのような流れですか？",
             "A: 以下のスケジュールを予定しています。\n9:00 受付集合、人事担当がお迎え\n9:15 社長挨拶、会社説明\n10:00 人事手続き（雇用契約書、誓約書等の記入）\n11:00 PC・社用スマホ等の貸与、セットアップ\n12:00 昼食（先輩社員とランチ）\n13:00 社内見学、各部署紹介\n14:00 配属部署へ、上長・チームメンバーとの顔合わせ\n15:00 OJT開始（業務説明）\n18:00 終業、歓迎会（任意参加）"),

            ("Q5: 初日の昼食はどうすればいいですか？",
             "A: 人事部が手配しますのでご安心ください。先輩社員と一緒に社員食堂またはランチ補助制度を利用した外食となります。費用は会社負担です。"),

            ("Q6: 名刺はいつもらえますか？",
             "A: 入社2週間後を目処に配布します。それまでは上長や先輩社員が同行しますので、名刺がなくても問題ありません。"),
        ]),

        ("3. 勤怠・給与関連", [
            ("Q7: 勤怠管理はどのように行いますか？",
             "A: 社内ポータルの勤怠管理システムで打刻します。出勤時・退勤時にシステムにログインし、「出勤」「退勤」ボタンをクリックしてください。リモートワーク時も同様です。"),

            ("Q8: 給与の支払日はいつですか？",
             "A: 毎月25日（休日の場合は前営業日）に支払われます。初回給与は入社月の翌月25日となります。例: 4月1日入社の場合、初回給与は5月25日。"),

            ("Q9: 残業代は出ますか？",
             "A: はい、出ます。時間外労働（1日8時間、週40時間を超える労働）には割増賃金（125%）が支払われます。管理監督者を除く全従業員が対象です。"),

            ("Q10: 有給休暇はいつから取得できますか？",
             "A: 入社6ヶ月後から10日間付与されます。それまでは有給休暇はありませんが、病気休暇（年10日、有給）を利用できます。"),
        ]),

        ("4. 福利厚生", [
            ("Q11: 社会保険はいつから加入できますか？",
             "A: 入社日から加入となります。健康保険証は入社2週間後にお渡しします。それまでに医療機関を受診する場合は、全額自己負担となりますが、後日保険証提示で払い戻しが受けられます。"),

            ("Q12: 通勤手当は支給されますか？",
             "A: はい。実費支給（上限月3万円）です。公共交通機関の場合は6ヶ月定期代の1/6を毎月支給します。自家用車通勤の場合は距離に応じたガソリン代相当額を支給します。"),

            ("Q13: 在宅勤務はできますか？",
             "A: はい。週3日まで在宅勤務可能です。ただし、入社後3ヶ月は原則出社し、業務を習得してから在宅勤務を開始することを推奨します。"),

            ("Q14: 社員食堂はありますか？",
             "A: 社員食堂はありませんが、ランチ補助制度があります。提携飲食店・コンビニで利用できる食事券を1食500円分、月20回まで支給します。"),
        ]),

        ("5. 社内システム", [
            ("Q15: 社内ポータルのログインID・パスワードはいつもらえますか？",
             "A: 初日にIT部門から発行されます。メールアドレス、勤怠管理システム、経費精算システム等、全てのシステムに同じID・パスワードでログインできます。"),

            ("Q16: 個人用メールアドレスはどうなりますか？",
             "A: 初日に付与されます。形式は「名前.苗字@tech-innovation.co.jp」です。例: 山田太郎さんの場合、「taro.yamada@tech-innovation.co.jp」"),

            ("Q17: SlackやTeamsは使いますか？",
             "A: 当社ではSlackを主要なコミュニケーションツールとして使用しています。初日にアカウントを発行し、使い方を説明します。"),

            ("Q18: 社内Wi-Fiのパスワードは？",
             "A: SSID「TechInnovation-Staff」、パスワードは初日にIT部門からお伝えします。社外秘情報ですので、外部の方には絶対に教えないでください。"),
        ]),

        ("6. オフィス環境", [
            ("Q19: 座席はフリーアドレスですか？",
             "A: いいえ。固定席制です。初日に自分の席が決まっており、デスク・椅子・PC等が用意されています。"),

            ("Q20: ロッカーはありますか？",
             "A: はい。各フロアに個人用ロッカーがあります。初日に鍵をお渡しします。私物の保管にご利用ください。"),

            ("Q21: 喫煙所はありますか？",
             "A: 社内は全面禁煙です。喫煙される方は、ビル1階の屋外喫煙所をご利用ください。"),

            ("Q22: 休憩スペースはありますか？",
             "A: 3階にカフェスペース、4階に休憩室があります。コーヒー・紅茶・お菓子は無料でご利用いただけます。"),
        ]),

        ("7. その他", [
            ("Q23: 研修制度はありますか？",
             "A: はい。入社後1ヶ月間は新入社員研修を実施します。ビジネスマナー、社内システムの使い方、業務知識等を学びます。その後、OJTで実務を習得します。"),

            ("Q24: 資格取得の支援はありますか？",
             "A: はい。業務に関連する資格の受験料は会社が全額補助します。また、合格時には合格祝い金（1万円〜5万円）が支給されます。"),

            ("Q25: 相談窓口はありますか？",
             "A: 人事部が窓口となります。業務上の悩み、人間関係、ハラスメント等、何でもご相談ください。また、外部カウンセラー（無料、匿名可）もご利用いただけます。"),

            ("Q26: 試用期間はありますか？",
             "A: はい。入社後3ヶ月間は試用期間です。この期間中の勤務成績が良好であれば、本採用となります。試用期間中も給与・待遇は本採用時と同じです。"),

            ("Q27: 服装・髪型に規定はありますか？",
             "A: 特に厳しい規定はありません。ビジネスカジュアルで、清潔感のある服装・髪型であれば問題ありません。来客対応時はビジネスフォーマルを推奨します。"),

            ("Q28: 副業は可能ですか？",
             "A: 可能です。ただし、事前申請・許可制です。本業に支障をきたさない範囲で、競合他社でない場合に限り許可されます。"),
        ]),
    ]

    for section_title, questions in faqs:
        doc.add_heading(section_title, 1)
        for q, a in questions:
            # 質問
            q_para = doc.add_paragraph()
            q_run = q_para.add_run(q)
            q_run.bold = True
            q_run.font.size = Pt(11)
            q_run.font.color.rgb = RGBColor(0, 102, 204)

            # 回答
            a_para = doc.add_paragraph(a)
            a_para.paragraph_format.left_indent = Inches(0.3)
            doc.add_paragraph()

    # お問い合わせ
    doc.add_page_break()
    doc.add_heading('お問い合わせ', 1)
    contact = doc.add_paragraph()
    contact.add_run('人事部 新入社員担当\n').bold = True
    contact.add_run('メール: jinji@tech-innovation.co.jp\n')
    contact.add_run('内線: 1234\n')
    contact.add_run('受付時間: 平日 9:00〜18:00')

    doc.save('data/raw/新入社員FAQ.docx')
    print("✓ 新入社員FAQ.docx を作成しました")

def create_it_faq_docx():
    """IT関連FAQ DOCX作成"""
    doc = Document()

    title = doc.add_heading('IT関連FAQ', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph('株式会社テックイノベーション IT部門')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph('最終更新: 2024年10月1日')
    doc.add_paragraph()

    doc.add_heading('目次', 1)
    doc.add_paragraph('1. アカウント・パスワード')
    doc.add_paragraph('2. メール')
    doc.add_paragraph('3. ネットワーク・VPN')
    doc.add_paragraph('4. PC・周辺機器')
    doc.add_paragraph('5. ソフトウェア')
    doc.add_paragraph('6. セキュリティ')
    doc.add_page_break()

    faqs = [
        ("1. アカウント・パスワード", [
            ("Q1: パスワードを忘れました。どうすればいいですか？",
             "A: 社内ポータルのログイン画面で「パスワードを忘れた場合」をクリックし、登録済みのメールアドレスに再設定用のリンクを送信してください。メールアドレスも不明な場合は、IT部門（内線1300）にお問い合わせください。本人確認後、パスワードをリセットします。"),

            ("Q2: パスワードの変更方法は？",
             "A: 社内ポータル → 「設定」 → 「パスワード変更」から変更できます。パスワードは3ヶ月ごとに変更が必要です。期限が近づくと通知メールが届きます。"),

            ("Q3: パスワードの要件は？",
             "A: 以下の要件を満たす必要があります。\n・8文字以上\n・英大文字、英小文字、数字、記号をそれぞれ1文字以上含む\n・過去3回のパスワードと異なること\n・ユーザー名、生年月日等の推測されやすい文字列を含まないこと"),
        ]),

        ("2. メール", [
            ("Q4: メールの容量上限は？",
             "A: メールボックスの容量上限は50GBです。上限に達すると新規メールを受信できなくなるため、定期的に不要メールを削除してください。現在の使用量は、メールクライアントの「容量確認」で確認できます。"),

            ("Q5: 大容量ファイルを送信したい場合は？",
             "A: メール添付は10MBまでです。それ以上のファイルは、社内ファイル共有システム（SharePoint）にアップロードし、共有リンクをメールで送信してください。外部宛の場合は、ファイル転送サービス（GigaFile便等）をご利用ください。"),

            ("Q6: 誤送信してしまいました。取り消せますか？",
             "A: 社内宛メールであれば、送信後1分以内なら取り消せます。Outlookの「送信済みアイテム」から該当メールを開き、「このメッセージを取り消す」をクリックしてください。社外宛メールは取り消せませんので、送信前に宛先・内容を十分確認してください。"),

            ("Q7: 迷惑メールが大量に届きます。",
             "A: 迷惑メールフィルタを強化します。該当メールを「迷惑メール」フォルダに移動するか、IT部門にメールを転送してください。送信者をブロックリストに追加します。"),
        ]),

        ("3. ネットワーク・VPN", [
            ("Q8: 社内Wi-Fiに接続できません。",
             "A: 以下を確認してください。\n1. SSID「TechInnovation-Staff」を選択しているか\n2. パスワードが正しいか（IT部門に確認）\n3. Wi-Fiがオンになっているか\n4. 機内モードがオフになっているか\nそれでも接続できない場合は、IT部門（内線1300）にお問い合わせください。"),

            ("Q9: VPN接続方法は？",
             "A: 以下の手順で接続してください。\n1. デスクトップのVPNアイコンをダブルクリック\n2. ユーザー名とパスワードを入力（社内アカウントと同じ）\n3. 「接続」をクリック\n4. 接続成功のメッセージが表示されれば完了\n詳細は社内ポータルの「VPN接続マニュアル」をご参照ください。"),

            ("Q10: VPNに接続できません。",
             "A: 以下を確認してください。\n1. インターネットに接続しているか\n2. ユーザー名・パスワードが正しいか\n3. VPNソフトが最新版か\n4. ファイアウォールがVPN通信を許可しているか\n解決しない場合は、IT部門にお問い合わせください。"),

            ("Q11: 外出先でインターネットに接続するには？",
             "A: 会社貸与のモバイルWi-Fiルーター、またはスマートフォンのテザリングをご利用ください。公衆Wi-Fiを使用する場合は、必ずVPN接続してから業務を行ってください。VPN未接続での機密情報の送受信は禁止です。"),
        ]),

        ("4. PC・周辺機器", [
            ("Q12: PCが起動しません。",
             "A: 以下を試してください。\n1. 電源ケーブルが接続されているか確認\n2. バッテリーが充電されているか確認\n3. 電源ボタンを10秒長押しして強制再起動\n4. 外付け機器を全て外して起動\nそれでも起動しない場合は、IT部門にお持ちください。代替機を貸与します。"),

            ("Q13: PCの動作が遅いです。",
             "A: 以下を試してください。\n1. 不要なアプリケーションを終了\n2. PCを再起動\n3. ディスク容量を確認（Cドライブの空き容量が10GB以上あるか）\n4. ウイルススキャンを実行\n改善しない場合は、IT部門にご相談ください。メモリ増設や買い替えを検討します。"),

            ("Q14: プリンターで印刷できません。",
             "A: 以下を確認してください。\n1. プリンターの電源が入っているか\n2. PCとプリンターが同じネットワークに接続されているか\n3. 用紙切れ、インク切れではないか\n4. 印刷ジョブが溜まっていないか（印刷キューを確認）\n解決しない場合は、IT部門にお問い合わせください。"),

            ("Q15: モニターが映りません。",
             "A: 以下を確認してください。\n1. モニターの電源が入っているか\n2. ケーブル（HDMI、USB-C）が正しく接続されているか\n3. PCの外部出力設定（Windows: Win+P、Mac: システム環境設定）\n4. 別のケーブル・ポートで試す\n解決しない場合は、IT部門にご連絡ください。"),
        ]),

        ("5. ソフトウェア", [
            ("Q16: 新しいソフトウェアをインストールしたいです。",
             "A: 管理者権限が必要なため、IT部門に申請してください。社内ポータル → 「IT関連」 → 「ソフトウェアインストール申請」から申請できます。セキュリティ審査後、許可されればIT部門がインストールします。"),

            ("Q17: Microsoft Officeが使えません。",
             "A: ライセンス認証が必要です。Word、Excel等を起動し、サインインを求められたら、会社のメールアドレスとパスワードを入力してください。それでも使えない場合は、IT部門にお問い合わせください。"),

            ("Q18: Zoomが使えません。",
             "A: 以下を確認してください。\n1. Zoomアプリが最新版か（アップデートを実行）\n2. カメラ・マイクの使用許可が有効か\n3. インターネット接続が安定しているか\n社内会議の場合は、Microsoft Teamsもご利用いただけます。"),
        ]),

        ("6. セキュリティ", [
            ("Q19: フィッシングメールが届きました。",
             "A: 絶対にリンクをクリックしたり、添付ファイルを開いたりしないでください。直ちにIT部門（security@tech-innovation.co.jp）にメールを転送し、その後削除してください。"),

            ("Q20: PCがウイルスに感染した疑いがあります。",
             "A: 直ちに以下を実施してください。\n1. ネットワークから切断（Wi-FiオフまたはLANケーブルを抜く）\n2. PCをシャットダウン\n3. IT部門に緊急連絡（内線1300、時間外: 080-XXXX-XXXX）\nIT部門が駆けつけ、対応します。他のPCは使用せず、感染拡大を防いでください。"),

            ("Q21: PCを紛失・盗難されました。",
             "A: 直ちにIT部門（080-XXXX-XXXX、24時間対応）に連絡してください。リモートでデータを消去します。警察に届け出た場合は、受理番号もIT部門にお伝えください。"),

            ("Q22: USBメモリを使用してもいいですか？",
             "A: 会社貸与のセキュリティUSBのみ使用可能です。個人のUSBメモリは使用禁止です。データの受け渡しは、社内ファイル共有システム（SharePoint）またはメール添付をご利用ください。"),
        ]),
    ]

    for section_title, questions in faqs:
        doc.add_heading(section_title, 1)
        for q, a in questions:
            q_para = doc.add_paragraph()
            q_run = q_para.add_run(q)
            q_run.bold = True
            q_run.font.size = Pt(11)
            q_run.font.color.rgb = RGBColor(204, 0, 0)

            a_para = doc.add_paragraph(a)
            a_para.paragraph_format.left_indent = Inches(0.3)
            doc.add_paragraph()

    doc.add_page_break()
    doc.add_heading('お問い合わせ', 1)
    contact = doc.add_paragraph()
    contact.add_run('IT部門 ヘルプデスク\n').bold = True
    contact.add_run('メール: it-support@tech-innovation.co.jp\n')
    contact.add_run('緊急連絡先: security@tech-innovation.co.jp\n')
    contact.add_run('内線: 1300\n')
    contact.add_run('緊急時(24時間): 080-XXXX-XXXX\n')
    contact.add_run('受付時間: 平日 9:00〜18:00')

    doc.save('data/raw/IT関連FAQ.docx')
    print("✓ IT関連FAQ.docx を作成しました")

def create_soumu_faq_docx():
    """総務FAQ DOCX作成"""
    doc = Document()

    title = doc.add_heading('総務FAQ', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph('株式会社テックイノベーション 総務部')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph('最終更新: 2024年10月1日')
    doc.add_paragraph()

    doc.add_heading('目次', 1)
    doc.add_paragraph('1. オフィス利用')
    doc.add_paragraph('2. 駐車場・交通')
    doc.add_paragraph('3. 備品・消耗品')
    doc.add_paragraph('4. 郵便・宅配')
    doc.add_paragraph('5. 名刺')
    doc.add_paragraph('6. その他')
    doc.add_page_break()

    faqs = [
        ("1. オフィス利用", [
            ("Q1: オフィスの利用時間は？",
             "A: 平日7:00〜22:00です。この時間外の利用は原則禁止です。休日に出社する場合は、事前に総務部に申請してください。セキュリティカードで入退館を記録しています。"),

            ("Q2: フリードリンクは何がありますか？",
             "A: 3階カフェスペースに、コーヒー、紅茶、緑茶、ウーロン茶、水を用意しています。無料でご利用いただけます。利用後はマグカップを洗い、元の位置に戻してください。"),

            ("Q3: 来客の受付方法は？",
             "A: 1階受付にお客様がお越しになったら、内線で連絡が入ります。速やかに1階までお迎えに行き、会議室までご案内ください。"),

            ("Q4: ゴミの分別方法は？",
             "A: 各階に分別ゴミ箱を設置しています。\n・燃えるゴミ（紙、生ゴミ等）\n・プラスチック\n・缶・ビン\n・ペットボトル\n・機密文書（シュレッダーボックス）\nルールを守り、適切に分別してください。"),
        ]),

        ("2. 駐車場・交通", [
            ("Q5: 駐車場はありますか？",
             "A: はい。ビル地下に社員専用駐車場があります（有料: 月1万円）。利用希望者は総務部に申請してください。空きがある場合のみ利用可能です。"),

            ("Q6: 自転車・バイク通勤は可能ですか？",
             "A: 可能です。ビル1階に駐輪場があります（無料）。利用する場合は総務部に届出が必要です。駐輪ステッカーを発行しますので、自転車・バイクに貼付してください。"),

            ("Q7: 通勤経路を変更したい場合は？",
             "A: 社内ポータル → 「人事関連」 → 「通勤経路変更届」から申請してください。定期代の精算が必要な場合は、定期券のコピーを添付してください。"),

            ("Q8: 定期券の区間を変更したい場合は？",
             "A: 通勤経路変更届と同時に、定期券の払い戻し証明書を提出してください。新しい定期代との差額を精算します。"),
        ]),

        ("3. 備品・消耗品", [
            ("Q9: 文房具が欲しい場合は？",
             "A: 3階305号室（総務部備品倉庫）に常備しています。自由にお取りください。在庫が少なくなったら、総務部にお知らせください。"),

            ("Q10: 名刺の追加発注方法は？",
             "A: 社内ポータル → 「総務関連」 → 「名刺発注」から申請してください。100枚単位で発注可能です。初回は無料、2回目以降は500円/100枚です。納期は5営業日です。"),

            ("Q11: PCの周辺機器が欲しい場合は？",
             "A: IT部門に申請してください。キーボード、マウス、モニター等は、業務上必要と認められれば貸与されます。"),

            ("Q12: 引き出しの鍵を紛失しました。",
             "A: 総務部にご連絡ください。スペアキーで開錠し、新しい鍵をお渡しします（実費負担: 1,000円）。"),
        ]),

        ("4. 郵便・宅配", [
            ("Q13: 郵便物の発送方法は？",
             "A: 1階受付の郵便ボックスに投函してください。毎日15:00に集荷します。宅急便は受付に持参すれば、総務部が発送手配します。送料は会社負担です。"),

            ("Q14: 個人宛の荷物を会社で受け取れますか？",
             "A: 可能です。ただし、頻繁な個人利用は避けてください。受付で受け取り、速やかに持ち帰ってください。"),

            ("Q15: 書留・速達を出したい場合は？",
             "A: 受付にお持ちいただければ、総務部が郵便局に依頼します。書留・速達の費用は会社負担です。"),
        ]),

        ("5. 名刺", [
            ("Q16: 名刺のデザインを変更できますか？",
             "A: いいえ。会社指定のデザインに統一しています。ただし、部署名、役職、連絡先は個別に記載可能です。"),

            ("Q17: 英語の名刺は作れますか？",
             "A: はい。海外顧客との商談が多い方は、英語版名刺を発行できます。総務部にご相談ください。"),

            ("Q18: 名刺を切らしてしまいました。緊急で必要です。",
             "A: 総務部にご連絡ください。即日発行（10枚まで）も対応可能です。ただし、通常の名刺よりも品質が劣ります。"),
        ]),

        ("6. その他", [
            ("Q19: 社員証を紛失しました。",
             "A: 直ちに総務部にご連絡ください。不正利用防止のため、すぐに無効化します。再発行手数料3,000円で新しい社員証を発行します。"),

            ("Q20: オフィスの温度調整はどうすればいいですか？",
             "A: 各フロアにエアコンリモコンがあります。温度設定は夏季28℃、冬季20℃を基本とし、体感に応じて±2℃の範囲で調整してください。"),

            ("Q21: ロッカーの鍵を紛失しました。",
             "A: 総務部にご連絡ください。スペアキーで開錠し、新しい鍵をお渡しします（実費負担: 1,000円）。"),

            ("Q22: 忘れ物をしました。",
             "A: 1階受付の遺失物保管箱を確認してください。保管期間は1ヶ月です。それ以降は処分します。"),

            ("Q23: 社内イベントの開催方法は？",
             "A: 総務部にご相談ください。会議室予約、参加者募集、費用精算等をサポートします。部署懇親会は1人5,000円まで補助があります。"),

            ("Q24: 災害時の避難場所は？",
             "A: ビル前の広場が一時避難場所です。地震・火災発生時は、エレベーターを使わず、階段で避難してください。避難訓練は年2回実施します。"),

            ("Q25: AEDはどこにありますか？",
             "A: 各階のエレベーター横に設置しています。使用方法は音声ガイドに従ってください。救急車の手配も同時に行ってください（119番通報）。"),
        ]),
    ]

    for section_title, questions in faqs:
        doc.add_heading(section_title, 1)
        for q, a in questions:
            q_para = doc.add_paragraph()
            q_run = q_para.add_run(q)
            q_run.bold = True
            q_run.font.size = Pt(11)
            q_run.font.color.rgb = RGBColor(0, 153, 0)

            a_para = doc.add_paragraph(a)
            a_para.paragraph_format.left_indent = Inches(0.3)
            doc.add_paragraph()

    doc.add_page_break()
    doc.add_heading('お問い合わせ', 1)
    contact = doc.add_paragraph()
    contact.add_run('総務部\n').bold = True
    contact.add_run('メール: soumu@tech-innovation.co.jp\n')
    contact.add_run('内線: 1250\n')
    contact.add_run('受付時間: 平日 9:00〜18:00')

    doc.save('data/raw/総務FAQ.docx')
    print("✓ 総務FAQ.docx を作成しました")


# PDF作成関数（ReportLabの日本語対応改善版）
def create_pdf_from_docx(docx_file, pdf_file):
    """DOCXファイルの内容を読み取ってPDFを作成"""
    from docx import Document as DocxDocument

    doc = DocxDocument(docx_file)

    pdf_doc = SimpleDocTemplate(pdf_file, pagesize=A4,
                                rightMargin=20*mm, leftMargin=20*mm,
                                topMargin=20*mm, bottomMargin=20*mm)

    story = []
    styles = getSampleStyleSheet()

    # 日本語対応スタイル
    title_style = ParagraphStyle(
        'JapaneseTitle',
        parent=styles['Title'],
        fontName=japanese_font,
        fontSize=24,
        alignment=TA_CENTER,
        spaceAfter=20,
    )

    heading1_style = ParagraphStyle(
        'JapaneseHeading1',
        parent=styles['Heading1'],
        fontName=japanese_font,
        fontSize=16,
        spaceAfter=12,
        spaceBefore=12,
    )

    body_style = ParagraphStyle(
        'JapaneseBody',
        parent=styles['BodyText'],
        fontName=japanese_font,
        fontSize=10,
        leading=16,
    )

    # ドキュメントから段落を抽出してPDF化
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            story.append(Spacer(1, 6))
            continue

        # スタイル判定
        if para.style.name.startswith('Heading 1') or para.style.name == 'Title':
            if '目次' not in text and 'FAQ' in text:
                story.append(Paragraph(text, title_style))
            else:
                story.append(Paragraph(text, heading1_style))
        elif 'Q' in text[:5] and ':' in text:
            # 質問
            bold_style = ParagraphStyle('Bold', parent=body_style, fontName=japanese_font, textColor='blue')
            story.append(Paragraph(text, bold_style))
        else:
            story.append(Paragraph(text, body_style))

        story.append(Spacer(1, 6))

    pdf_doc.build(story)
    print(f"✓ {os.path.basename(pdf_file)} を作成しました")


if __name__ == '__main__':
    # DOCX作成
    create_shinnyusyain_faq_docx()
    create_it_faq_docx()
    create_soumu_faq_docx()

    # PDF作成
    print("\nPDFファイルを作成中...")
    create_pdf_from_docx('data/raw/新入社員FAQ.docx', 'data/raw/新入社員FAQ.pdf')
    create_pdf_from_docx('data/raw/IT関連FAQ.docx', 'data/raw/IT関連FAQ.pdf')
    create_pdf_from_docx('data/raw/総務FAQ.docx', 'data/raw/総務FAQ.pdf')

    print("\n✅ 全てのFAQファイル（DOCX・PDF）を作成しました！")
